from dataclasses import dataclass
import json, time
from typing import Dict, List
from docx import Document, shared

RED = shared.RGBColor(0x7D, 0x08, 0x00)
ORANGE = shared.RGBColor(0x66, 0x58, 0x00)
BLACK = shared.RGBColor(0x00, 0x00, 0x00)


def load(filepath) -> dict:
    with open(filepath, encoding='utf-8') as infile:
        return json.load(infile)



@dataclass
class Word:
    content: str
    confidence: float


def get_words_by_start_time(transcript: Dict) -> Dict[str, Word]:
    """Merges punctuation with standard words since they don't have a start time,
    returns them in a handy map of start_time to word and confidence.

    Args:
        transcript: Amazon Transcript JSON dictionary

    Returns:
        (Dict): a map of start_time to word and confidence
    """
    merged_words: Dict[str, Word] = {}

    items: List[Dict] = transcript["results"]["items"]

    for i, item in enumerate(items):
        # Only save pronunciations... may not be necessary, or may need other types
        if item["type"] == "pronunciation":
            word: str = item["alternatives"][0]["content"]
            confidence: str = item["alternatives"][0]["confidence"]

            # If the next item in the transcript is a punctuation, merge with the current word
            if i < len(items) - 1 and items[i + 1]["type"] == "punctuation":
                word += items[i + 1]["alternatives"][0]["content"]

            # Add the word to the map at start time
            merged_words[item["start_time"]] = Word(word, float(confidence))

    return merged_words


def build_docx(title, word_time_map: Dict[str, Word], speaker_names: Dict[str, str]):
    doc = Document()
    doc.add_heading(title, 0)

    current_speaker: str = ""
    current_paragraph = None

    for speaker_segment in transcript["results"]["speaker_labels"]["segments"]:
        name = speaker_names[speaker_segment["speaker_label"]]

        if current_speaker != name:
            speaker_start = speaker_segment["start_time"]
            start = time.strftime("%Hh%Mm%Ss", time.gmtime(float(speaker_start)))

            doc.add_heading(f"{name} @ {start}", 2)
            current_speaker = name
            current_paragraph = doc.add_paragraph()

            if speaker_start in word_time_map:
                word_time_map[speaker_start].content = word_time_map[speaker_start].content.capitalize()

        for item in speaker_segment["items"]:
            start = item["start_time"]
            if start in word_time_map:
                word = word_time_map[start].content + " "
                current_paragraph.add_run(word).font.color.rgb = (
                    RED if word_time_map[start].confidence < 0.5
                    else ORANGE if word_time_map[start].confidence < 0.85
                    else BLACK
                )

    return doc



if __name__ == "__main__":
    # Set Amazon Transcribe transcript path
    transcript = load("./asrOutput.json")

    # Set desired output title
    title = "Document Title"

    # Set the mapping of speakers in the transcription
    name_map = {"spk_0": "Speaker One", "spk_1": "Speaker Two"}  # ... }

    words_by_time = get_words_by_start_time(transcript)
    document = build_docx(title, words_by_time, name_map)
    document.save(f"{title}.docx")
